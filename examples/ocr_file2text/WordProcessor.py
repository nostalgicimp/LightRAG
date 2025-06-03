import docx # 用于处理 .docx 文件
from docx.opc.constants import RELATIONSHIP_TYPE # 用于识别图像关系
from docx.document import Document as DocxDocument # 类型提示
# from docx.table import Table as DocxTable # 类型提示 (如果特定方法需要)
from typing import List, Tuple, Optional, Dict, Any
# import re # 当前版本未使用，可移除
from PIL import Image # 导入Pillow库
import pytesseract    # 导入pytesseract
import io             # 导入io用于处理字节流
import numpy as np    # PaddleOCR 需要

# 尝试导入PaddleOCR，并处理未安装的情况
try:
    from paddleocr import PaddleOCR
    PADDLEOCR_AVAILABLE = True  # PaddleOCR可用标志
except ImportError:
    PADDLEOCR_AVAILABLE = False # PaddleOCR不可用标志
    PaddleOCR = None # 定义PaddleOCR为None，以避免后续代码出现NameError

class WordProcessor:
    """使用 python-docx 实现的 Word (.docx) 文件处理器，支持Tesseract和PaddleOCR。"""

    def __init__(self,
                 tesseract_cmd_path: Optional[str] = None,
                 default_ocr_engine: str = 'paddle', # 默认OCR引擎
                 paddle_ocr_config: Optional[Dict[str, Any]] = None):
        """
        初始化 WordProcessor

        Args:
            tesseract_cmd_path (Optional[str]): Tesseract OCR的可执行文件路径。
                                                如果Tesseract在系统PATH中，则无需设置。
            default_ocr_engine (str): 默认使用的OCR引擎 ('tesseract' 或 'paddle')。
            paddle_ocr_config (Optional[Dict[str, Any]]): PaddleOCR的初始化配置。
                例如: {'use_angle_cls': True, 'lang': 'ch', 'show_log': False}
        """
        self.tesseract_available = False
        if tesseract_cmd_path:
            pytesseract.pytesseract.tesseract_cmd = tesseract_cmd_path
        
        try:
            pytesseract.get_tesseract_version()
            self.tesseract_available = True
        except pytesseract.TesseractNotFoundError:
            print("警告: Tesseract OCR 未找到或未正确配置。如果选择Tesseract，图像文本提取功能将不可用。")
        except Exception as e:
            print(f"警告: 初始化Tesseract时发生错误: {e}。Tesseract图像文本提取功能可能受影响。")

        self.default_ocr_engine = default_ocr_engine.lower()
        if self.default_ocr_engine == 'paddle' and not PADDLEOCR_AVAILABLE:
            print("警告: PaddleOCR库未安装，但被选为默认OCR引擎。将尝试回退到Tesseract (如果可用)，否则不进行图像OCR。")
            if self.tesseract_available:
                self.default_ocr_engine = 'tesseract'
            else:
                self.default_ocr_engine = None

        self.paddle_ocr_instance: Optional[PaddleOCR] = None
        self.user_paddle_ocr_config = paddle_ocr_config
        self.current_paddle_lang_init = None

    def _initialize_paddleocr(self, lang_for_paddle: str) -> bool:
        """【私有方法】惰性初始化PaddleOCR实例。"""
        if not PADDLEOCR_AVAILABLE:
            print("警告: PaddleOCR库未安装，无法使用PaddleOCR引擎。")
            return False

        if self.user_paddle_ocr_config:
            should_reinitialize_user_config = self.paddle_ocr_instance is None
            if 'lang' in self.user_paddle_ocr_config:
                 should_reinitialize_user_config = should_reinitialize_user_config or \
                    (self.current_paddle_lang_init != self.user_paddle_ocr_config.get('lang'))
            if should_reinitialize_user_config:
                try:
                    print(f"正在使用用户配置初始化PaddleOCR: {self.user_paddle_ocr_config}...")
                    self.paddle_ocr_instance = PaddleOCR(**self.user_paddle_ocr_config)
                    self.current_paddle_lang_init = self.user_paddle_ocr_config.get('lang', 'ch')
                    print("PaddleOCR已通过用户配置成功初始化。")
                    return True
                except Exception as e:
                    print(f"警告: 使用用户配置初始化PaddleOCR失败: {e}")
                    self.paddle_ocr_instance = None
                    return False
            return True

        if self.paddle_ocr_instance is None or self.current_paddle_lang_init != lang_for_paddle:
            try:
                default_config = {'use_angle_cls': True, 'lang': lang_for_paddle, 'show_log': False}
                print(f"正在使用默认配置初始化PaddleOCR (语言: {lang_for_paddle})...")
                self.paddle_ocr_instance = PaddleOCR(**default_config)
                self.current_paddle_lang_init = lang_for_paddle
                print("PaddleOCR已成功初始化。")
                return True
            except Exception as e:
                print(f"警告: 初始化PaddleOCR失败 (语言={lang_for_paddle}): {e}")
                self.paddle_ocr_instance = None
                return False
        return True

    def _map_tesseract_lang_to_paddle(self, tesseract_langs: str) -> str:
        """【私有方法】启发式地将Tesseract语言字符串映射到PaddleOCR语言代码。"""
        if self.user_paddle_ocr_config and 'lang' in self.user_paddle_ocr_config:
            return self.user_paddle_ocr_config['lang']
        t_langs = tesseract_langs.lower()
        if 'chi_sim' in t_langs or 'chi_tra' in t_langs: return 'ch'
        if 'eng' in t_langs: return 'en'
        print(f"警告: 无法将Tesseract语言 '{tesseract_langs}'直接映射到PaddleOCR支持的语言。将默认使用 'ch'。")
        return 'ch'

    def extract_text_from_word(self, word_path: str,
                               include_metadata: bool = False,
                               include_images: bool = False,
                               ocr_languages: str = 'chi_sim+eng',
                               ocr_engine: Optional[str] = None) -> Tuple[str, Optional[Dict[str, Any]]]:
        """
        从Word (.docx)文件中提取文本内容。

        Args:
            word_path (str): Word文件路径 (.docx)。
            include_metadata (bool): 是否包含元数据。
            include_images (bool): 是否尝试提取图像中的文本。
            ocr_languages (str): OCR识别时使用的语言。
            ocr_engine (Optional[str]): 使用的OCR引擎 ('tesseract', 'paddle')。如果None，使用默认引擎。

        Returns:
            tuple: (提取的文本, 元数据字典(如果include_metadata为True))
        """
        text_parts = []
        metadata_dict: Optional[Dict[str, Any]] = None
        chosen_ocr_engine = (ocr_engine or self.default_ocr_engine or "none").lower()

        try:
            doc = docx.Document(word_path)

            # 1. 提取段落文本
            for para in doc.paragraphs:
                text_parts.append(para.text)

            # 2. 提取表格中的文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = "\n".join([p.text for p in cell.paragraphs])
                        text_parts.append(cell_text) # strip later or join with space
            
            # 合并主要文本内容 (段落和表格)
            # 使用空格连接，因为段落和单元格内容本身可能需要换行
            main_content_text = " ".join(filter(None, text_parts)).strip()


            # 3. 提取元数据
            if include_metadata:
                cp = doc.core_properties
                metadata_dict = {
                    "作者": cp.author,
                    "类别": cp.category,
                    "备注": cp.comments,
                    "内容状态": cp.content_status,
                    "创建日期": cp.created.isoformat() if cp.created else None,
                    "标识符": cp.identifier,
                    "关键词": cp.keywords,
                    "语言": cp.language,
                    "最后修改者": cp.last_modified_by,
                    "最后打印日期": cp.last_printed.isoformat() if cp.last_printed else None,
                    "修改日期": cp.modified.isoformat() if cp.modified else None,
                    "修订号": cp.revision,
                    "主题": cp.subject,
                    "标题": cp.title,
                    "版本": cp.version
                }

            # 4. 提取图像文本
            image_ocr_text_parts = []
            if include_images:
                if chosen_ocr_engine == "none":
                    print("信息: 未选择OCR引擎 (或默认引擎不可用)，跳过图像文本提取。")
                elif (chosen_ocr_engine == 'tesseract' and not self.tesseract_available) or \
                     (chosen_ocr_engine == 'paddle' and not PADDLEOCR_AVAILABLE):
                    print(f"警告: 选择的OCR引擎 '{chosen_ocr_engine}' 不可用。跳过图像文本提取。")
                else:
                    try:
                        image_text_content = self._extract_text_from_word_images(
                            doc, ocr_languages, chosen_ocr_engine
                        )
                        if image_text_content.strip():
                            image_ocr_text_parts.append(f"\n--- 来自图像的文本 ({chosen_ocr_engine}) ---\n{image_text_content.strip()}")
                    except Exception as e_ocr: # 更通用的异常捕获
                        print(f"警告: 从 {word_path} 的图像中提取文本时发生错误 ({chosen_ocr_engine}): {e_ocr}")
            
            # 组合所有文本
            final_text_elements = [main_content_text] + image_ocr_text_parts
            full_text_output = "\n".join(filter(None, final_text_elements)).strip()


        except docx.opc.exceptions.PackageNotFoundError:
             raise FileNotFoundError(f"错误: Word文件未找到或不是有效的 .docx 格式: {word_path}")
        except Exception as e:
            raise Exception(f"处理Word文件 '{word_path}' 时发生错误: {str(e)}")

        return full_text_output, metadata_dict

    def _ocr_image_with_tesseract_from_bytes(self, image_bytes: bytes, languages: str, image_id: str) -> str:
        """【私有方法】使用Tesseract从图像字节流OCR。"""
        try:
            pil_image = Image.open(io.BytesIO(image_bytes))
            if pil_image.mode == 'RGBA' or pil_image.mode == 'P':
                pil_image = pil_image.convert('RGB')
            text = pytesseract.image_to_string(pil_image, lang=languages)
            return text.strip()
        except pytesseract.TesseractError as te:
            print(f"  - Tesseract OCR错误 (图像ID: {image_id}): {te}")
        except Image.UnidentifiedImageError:
            print(f"  - Tesseract警告: 无法识别图像格式 (图像ID: {image_id})。")
        except Exception as e_img:
            print(f"  - Tesseract处理图像 (图像ID: {image_id}) 时发生未知错误: {e_img}")
        return ""

    def _ocr_image_with_paddle_from_bytes(self, image_bytes: bytes, lang_for_paddle: str, image_id: str) -> str:
        """【私有方法】使用PaddleOCR从图像字节流OCR。"""
        if not self._initialize_paddleocr(lang_for_paddle) or not self.paddle_ocr_instance:
            print(f"  - PaddleOCR初始化失败或不可用，无法处理图像 (图像ID: {image_id})")
            return ""
        try:
            pil_image = Image.open(io.BytesIO(image_bytes))
            if pil_image.mode == 'RGBA' or pil_image.mode == 'P':
                pil_image = pil_image.convert('RGB')
            elif pil_image.mode == 'L':
                 pil_image = pil_image.convert('RGB')
            
            img_np = np.array(pil_image)
            ocr_results = self.paddle_ocr_instance.ocr(img_np, cls=True)
            
            image_texts = []
            if ocr_results and ocr_results[0] is not None:
                for line_info in ocr_results[0]:
                    text_content = line_info[1][0]
                    if text_content.strip():
                        image_texts.append(text_content.strip())
            return "\n".join(image_texts)
        except Image.UnidentifiedImageError:
            print(f"  - PaddleOCR警告: 无法识别图像格式 (图像ID: {image_id})。")
        except Exception as e_img:
            print(f"  - PaddleOCR处理图像 (图像ID: {image_id}) 时发生错误: {e_img}")
        return ""

    def _extract_text_from_word_images(self, 
                                       doc: DocxDocument, 
                                       ocr_languages: str = 'chi_sim+eng',
                                       ocr_engine: str = 'tesseract') -> str:
        """
        【私有方法】从Word文档中的图像提取文本。

        Args:
            doc (DocxDocument): python-docx 的文档对象。
            ocr_languages (str): OCR识别时使用的语言。
            ocr_engine (str): 使用的OCR引擎。

        Returns:
            str: 从所有图像中提取并合并的文本。
        """
        image_text_parts = []
        image_count = 0
        
        print(f"OCR ({ocr_engine}): 正在扫描Word文档以查找图像...")
        # 图像可以嵌入在段落中 (InlineShape) 或作为文档的一部分 (rels)
        # 首先尝试通过rels，这是更通用的方式
        # 也要考虑 doc.inline_shapes 中的图片

        # 方法1: 通过文档关系查找 (通常能找到大部分图片)
        rels = doc.part.rels
        for rel_id in rels:
            try:
                rel = rels[rel_id]
                if rel.reltype == RELATIONSHIP_TYPE.IMAGE:
                    image_count += 1
                    image_part = rel.target_part
                    image_bytes = image_part.blob # 图像的字节数据
                    image_content_type = image_part.content_type # 例如 'image/png'
                    
                    # print(f"  - 找到图像 {image_count} (通过rels, 类型: {image_content_type}, ID: {rel_id})")
                    text_from_one_image = ""
                    if ocr_engine == 'tesseract':
                        text_from_one_image = self._ocr_image_with_tesseract_from_bytes(image_bytes, ocr_languages, rel_id)
                    elif ocr_engine == 'paddle':
                        lang_for_paddle = self._map_tesseract_lang_to_paddle(ocr_languages)
                        text_from_one_image = self._ocr_image_with_paddle_from_bytes(image_bytes, lang_for_paddle, rel_id)
                    
                    if text_from_one_image:
                        image_text_parts.append(f"\n--- 图像 {image_count} (ID: {rel_id}) ---\n{text_from_one_image}")
            except Exception as e_rel:
                print(f"  - 警告: 处理文档关系中的项 (ID: {rel_id}) 时出错: {e_rel}")
        
        # 方法2: 遍历内联形状 (InlineShapes)
        # 这可能会找到一些通过rels未直接获取的，或重复获取（需要注意）
        # python-docx的inline_shapes主要用于访问形状属性，获取图像数据可能仍需通过rels或blipFill
        # 这里简化处理，假设rels已覆盖大部分，或用户可扩展此部分
        # for inline_shape in doc.inline_shapes:
        #     if inline_shape.type == docx.enum.shape.WD_INLINE_SHAPE.PICTURE:
        #         # 获取图像数据比较复杂，可能需要深入解析XML (blip_rId = inline_shape.blip_rId)
        #         # 然后通过 blip_rId 在 doc.part.rels 中找到对应的图像
        #         # 为避免重复，如果主要通过rels，此部分可以简化或作为补充
        #         pass

        if image_count == 0:
            print(f"OCR ({ocr_engine}): 在文档中未通过关系(rels)找到可识别的图像。")
        else:
            print(f"OCR ({ocr_engine}): 完成扫描，共处理 {image_count} 个通过关系(rels)找到的图像。")
            
        return "\n".join(image_text_parts)

    def extract_tables(self, word_path: str) -> List[List[List[str]]]:
        """
        提取Word (.docx)文件中的所有表格数据。

        Args:
            word_path (str): Word文件路径 (.docx)。

        Returns:
            List[List[List[str]]]: 三维列表，结构为 [表格索引][行索引][单元格文本]。
        """
        all_tables_extracted: List[List[List[str]]] = []
        try:
            doc = docx.Document(word_path)
            if doc.tables:
                print(f"表格提取: 在文档中找到 {len(doc.tables)} 个表格。")
                for table_idx, table_obj in enumerate(doc.tables):
                    current_table_data: List[List[str]] = []
                    for row_idx, row in enumerate(table_obj.rows):
                        current_row_data: List[str] = []
                        for cell_idx, cell in enumerate(row.cells):
                            # 合并单元格内所有段落的文本，段落间用换行符分隔
                            cell_text = "\n".join([p.text for p in cell.paragraphs])
                            current_row_data.append(cell_text.strip())
                        current_table_data.append(current_row_data)
                    all_tables_extracted.append(current_table_data)
                    # print(f"  - 表格 {table_idx + 1}: 提取了 {len(current_table_data)} 行。")
            else:
                print("表格提取: 在文档中未找到表格。")

        except docx.opc.exceptions.PackageNotFoundError:
             raise FileNotFoundError(f"错误: Word文件未找到或不是有效的 .docx 格式: {word_path}")
        except Exception as e:
            raise Exception(f"从Word文件提取表格时发生错误 ({word_path}): {str(e)}")
        return all_tables_extracted